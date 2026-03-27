"""
Document loader for procurement files.

Supports PDF, DOCX, TXT, and Markdown files commonly found in
government procurement workflows (RFPs, bid packages, policy memos).
"""

import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
import hashlib
import logging

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Represents a loaded document with metadata."""

    content: str
    source: str
    doc_type: str
    metadata: dict = field(default_factory=dict)

    @property
    def doc_id(self) -> str:
        """Deterministic document ID based on content hash."""
        return hashlib.sha256(self.content.encode()).hexdigest()[:16]

    @property
    def char_count(self) -> int:
        return len(self.content)

    @property
    def word_count(self) -> int:
        return len(self.content.split())


class DocumentLoader:
    """
    Multi-format document loader for procurement files.

    Handles the variety of file formats typically found in
    government procurement packages.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv"}

    def load_file(self, file_path: str) -> Document:
        """Load a single file and return a Document."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        loader_map = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".txt": self._load_text,
            ".md": self._load_text,
            ".csv": self._load_csv,
        }

        content = loader_map[ext](path)

        return Document(
            content=content,
            source=str(path.name),
            doc_type=ext.lstrip("."),
            metadata={
                "file_path": str(path.absolute()),
                "file_size_bytes": path.stat().st_size,
                "extension": ext,
            },
        )

    def load_directory(
        self, dir_path: str, recursive: bool = False
    ) -> list[Document]:
        """Load all supported documents from a directory."""
        path = Path(dir_path)
        if not path.is_dir():
            raise NotADirectoryError(f"Not a directory: {dir_path}")

        pattern = "**/*" if recursive else "*"
        documents = []

        for file_path in sorted(path.glob(pattern)):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.load_file(str(file_path))
                    documents.append(doc)
                    logger.info(
                        f"Loaded: {file_path.name} "
                        f"({doc.word_count} words)"
                    )
                except Exception as e:
                    logger.warning(f"Failed to load {file_path}: {e}")

        logger.info(
            f"Loaded {len(documents)} documents from {dir_path}"
        )
        return documents

    def _load_pdf(self, path: Path) -> str:
        """Extract text from PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError(
                "PyMuPDF required for PDF support: pip install PyMuPDF"
            )

        doc = fitz.open(str(path))
        pages = []
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                pages.append(f"[Page {page_num + 1}]\n{text}")
        doc.close()
        return "\n\n".join(pages)

    def _load_docx(self, path: Path) -> str:
        """Extract text from Word documents."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx required for DOCX support: pip install python-docx"
            )

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _load_text(self, path: Path) -> str:
        """Load plain text or markdown files."""
        return path.read_text(encoding="utf-8")

    def _load_csv(self, path: Path) -> str:
        """Load CSV as text representation for embedding."""
        import csv

        rows = []
        with open(path, newline="", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))
        return "\n".join(rows)
